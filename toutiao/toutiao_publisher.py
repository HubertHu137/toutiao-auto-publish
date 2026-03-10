#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
今日头条热榜新闻自动发布脚本 v2.1
流程：清理旧文档 -> 抓取热榜 -> 扩展过滤 -> 历史去重(70%相似度) -> 搜索原文+图片 -> AI改写 -> 保存图文MD -> 发布
新功能：
  - 自动清理7日以上的旧文档（DOCX、MD、JSON、图片）
  - 标题相似度阈值提升至70%（更严格的去重）
  - 每篇文章至少1张图片，不足则备用搜索
  - 每次生成5-6篇文章
"""

import sys
import json
import time
import re
import urllib.request
import urllib.error
import urllib.parse
import os
import glob
from datetime import datetime
from pathlib import Path

# Word文档生成
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx 未安装，安装中...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True

# ========== 配置 ==========
AI_API_URL  = "http://127.0.0.1:19999"
AI_API_KEY  = os.environ.get("AI_API_KEY", "")   # 必填：AI API Key，通过环境变量 AI_API_KEY 配置
TAVILY_API_KEY  = os.environ.get("TAVILY_API_KEY", "")   # 必填：Tavily Search API Key，通过环境变量 TAVILY_API_KEY 配置
MODEL_ID        = "gpt-4o-mini"
TARGET_MIN      = 5        # 目标最少文章数
TARGET_MAX      = 6        # 目标最多文章数（不超过6篇）
FETCH_POOL      = 40       # 从热榜抓取的候选池大小（多抓一些以备过滤）
MIN_IMAGES      = 1        # 每篇文章至少需要的图片数量
OUTPUT_DIR      = os.path.expanduser("~/.openclaw/toutiao_articles")
MD_DIR          = os.path.expanduser("~/.openclaw/toutiao_articles/md")  # MD 文章存放目录
DOCX_DIR        = os.path.expanduser("~/.openclaw/toutiao_articles/docx")  # Word 文档存放目录
HISTORY_FILE    = os.path.expanduser("~/.openclaw/toutiao_articles/history_titles.txt")  # 历史标题去重

# ========== 过滤关键词（政府/政治/经济政策类全覆盖）==========
FILTER_KEYWORDS = [
    # 领导人
    "习近平", "李强", "李克强", "王沪宁", "赵乐际", "丁薛祥", "李希",
    "韩正", "汪洋", "栗战书",
    # 机构/党政
    "中央", "政府", "人大", "政协", "党委", "国务院", "中共", "中纪委",
    "省委", "市委", "纪委", "监委", "政治局", "中办", "国办",
    # 职位
    "书记", "主席", "副主席", "总理", "副总理", "部长", "省长", "市长",
    "官员", "领导", "干部", "党员", "党建", "代表建议", "全国人大代表",
    # 媒体
    "人民日报", "新华社", "央视", "央广", "光明日报", "官方", "官媒",
    # 政策/法规
    "两会", "全国", "反腐", "纪律", "检察", "法院", "公检法",
    "公安部", "外交部", "国防部", "国防费", "军委", "解放军", "武警",
    "国防预算", "预算报告", "工作报告", "五年规划", "政府工作",
    # 外交/国际政治
    "台湾问题", "南海", "钓鱼岛", "制裁", "外交", "峰会", "联合国",
    # 经济政策（宏观政策类过滤，保留民生/消费类）
    "GDP目标", "经济预期目标", "货币政策", "财政政策", "赤字率",
    "专项债", "国债", "宏观调控",
    # 香港/澳门政治
    "行政长官", "特区政府", "立法会",
]

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(MD_DIR, exist_ok=True)
os.makedirs(DOCX_DIR, exist_ok=True)
Path(HISTORY_FILE).touch(exist_ok=True)


def log(msg):
    print(f"[{datetime.now().strftime('%H:%M:%S')}] {msg}", flush=True)


# ========== 工具：清理旧文档 ==========
def cleanup_old_documents(days=7):
    """
    清理7日以上的Word文档和相关文件
    保留：DOCX、MD、JSON文件
    """
    from datetime import timedelta
    
    log(f"🧹 清理 {days} 日前的旧文档...")
    
    cutoff_date = datetime.now() - timedelta(days=days)
    cutoff_timestamp = cutoff_date.timestamp()
    
    deleted_count = 0
    total_size = 0
    
    # 清理DOCX文件
    for docx_file in glob.glob(os.path.join(DOCX_DIR, "*.docx")):
        try:
            mtime = os.path.getmtime(docx_file)
            if mtime < cutoff_timestamp:
                size = os.path.getsize(docx_file)
                os.remove(docx_file)
                deleted_count += 1
                total_size += size
                log(f"  删除: {os.path.basename(docx_file)}")
        except Exception as e:
            log(f"  ⚠️  删除失败 {os.path.basename(docx_file)}: {e}")
    
    # 清理MD文件
    for md_file in glob.glob(os.path.join(MD_DIR, "*.md")):
        try:
            mtime = os.path.getmtime(md_file)
            if mtime < cutoff_timestamp:
                size = os.path.getsize(md_file)
                os.remove(md_file)
                deleted_count += 1
                total_size += size
        except Exception as e:
            pass
    
    # 清理JSON文件
    for json_file in glob.glob(os.path.join(OUTPUT_DIR, "*.json")):
        try:
            mtime = os.path.getmtime(json_file)
            if mtime < cutoff_timestamp:
                size = os.path.getsize(json_file)
                os.remove(json_file)
                deleted_count += 1
                total_size += size
        except Exception as e:
            pass
    
    # 清理images目录下的旧图片
    images_dir = os.path.join(OUTPUT_DIR, "images")
    if os.path.exists(images_dir):
        for img_file in glob.glob(os.path.join(images_dir, "*")):
            try:
                mtime = os.path.getmtime(img_file)
                if mtime < cutoff_timestamp:
                    size = os.path.getsize(img_file)
                    os.remove(img_file)
                    deleted_count += 1
                    total_size += size
            except Exception as e:
                pass
    
    if deleted_count > 0:
        size_mb = total_size / (1024 * 1024)
        log(f"  ✅ 已清理 {deleted_count} 个文件，释放空间 {size_mb:.2f} MB")
    else:
        log(f"  ✅ 无需清理，所有文件都在 {days} 天内")


# ========== 工具：加载历史标题 ==========
def load_history_titles():
    """加载近 30 天的历史标题（用于去重）"""
    try:
        with open(HISTORY_FILE, encoding="utf-8") as f:
            return [line.strip() for line in f if line.strip()]
    except Exception:
        return []


def save_history_title(title):
    """追加标题到历史记录"""
    with open(HISTORY_FILE, "a", encoding="utf-8") as f:
        f.write(title + "\n")


def is_similar_to_history(title, history_titles, threshold=0.7):
    """
    判断标题是否与历史标题相似（Jaccard 相似度）
    threshold=0.7 表示 70% 词汇重叠即视为相似
    """
    def tokenize(t):
        # 中文按字切，英文按词切
        chars = re.findall(r'[\u4e00-\u9fff]', t)
        words = re.findall(r'[a-zA-Z0-9]+', t.lower())
        return set(chars + words)

    title_tokens = tokenize(title)
    if not title_tokens:
        return False

    for hist in history_titles:
        hist_tokens = tokenize(hist)
        if not hist_tokens:
            continue
        inter = len(title_tokens & hist_tokens)
        union = len(title_tokens | hist_tokens)
        if union > 0 and inter / union >= threshold:
            return True
    return False


# ========== Step 1: 抓取头条热榜 ==========
def fetch_toutiao_hot(max_count=FETCH_POOL):
    """抓取今日头条热榜，返回尽量多的候选标题"""
    log("正在抓取今日头条热榜...")

    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                      "AppleWebKit/537.36 (KHTML, like Gecko) "
                      "Chrome/120.0.0.0 Safari/537.36",
        "Referer": "https://www.toutiao.com/",
        "Accept": "application/json, text/plain, */*",
        "Accept-Language": "zh-CN,zh;q=0.9",
    }

    titles = []
    raw_items = []  # 保留原始 item 数据（含图片URL）

    # 方式1：头条热榜 API
    for url in [
        "https://www.toutiao.com/hot-event/hot-board/?origin=toutiao_pc&_signature=",
        "https://api.toutiao.com/hot-event/hot-board/?origin=toutiao_pc",
    ]:
        try:
            req = urllib.request.Request(url, headers=headers)
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = json.loads(resp.read().decode("utf-8", errors="replace"))
            items = (data.get("data") or data.get("result", {}).get("data") or data.get("items") or [])
            for item in items[:max_count]:
                title = (item.get("Title") or item.get("title") or
                         item.get("query") or item.get("name") or "").strip()
                if title:
                    titles.append(title)
                    raw_items.append(item)
            if titles:
                log(f"  方式1成功，获取 {len(titles)} 条热榜")
                break
        except Exception as e:
            log(f"  方式1失败: {e}")

    # 方式2：微博热搜（稳定备用）
    if not titles:
        log("  尝试方式2：微博热搜...")
        try:
            url = "https://weibo.com/ajax/side/hotSearch"
            req = urllib.request.Request(url, headers=headers)
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = json.loads(resp.read().decode("utf-8", errors="replace"))
            items = data.get("data", {}).get("realtime", [])
            for item in items[:max_count]:
                title = item.get("word", "").strip()
                if title:
                    titles.append(title)
                    raw_items.append(item)
            if titles:
                log(f"  方式2成功，获取 {len(titles)} 条热搜")
        except Exception as e:
            log(f"  方式2失败: {e}")

    # 方式3：百度热搜备用
    if not titles:
        log("  尝试方式3：百度热搜...")
        try:
            url = "https://top.baidu.com/board?tab=realtime"
            req = urllib.request.Request(url, headers=headers)
            with urllib.request.urlopen(req, timeout=10) as resp:
                html = resp.read().decode("utf-8", errors="replace")
            matches = re.findall(
                r'class="c-single-text-ellipsis"[^>]*>([^<]+)<', html)
            if not matches:
                matches = re.findall(r'"title":"([^"]{4,30})"', html)
            for m in matches[:max_count]:
                title = m.strip()
                if title:
                    titles.append(title)
                    raw_items.append({"title": title})
            if titles:
                log(f"  方式3成功，获取 {len(titles)} 条热搜")
        except Exception as e:
            log(f"  方式3失败: {e}")

    return titles, raw_items


# ========== Step 2: 过滤 ==========
def filter_titles(titles, history_titles):
    """过滤政府/政治类 + 历史重复标题"""
    passed = []
    removed_gov = []
    removed_dup = []

    for title in titles:
        # 政府/政治关键词过滤
        if any(kw in title for kw in FILTER_KEYWORDS):
            removed_gov.append(title)
            continue
        # 历史去重
        if is_similar_to_history(title, history_titles):
            removed_dup.append(title)
            continue
        passed.append(title)

    log(f"  过滤政府/政策类：{len(removed_gov)} 条")
    if removed_gov:
        log(f"    {removed_gov[:5]}")
    log(f"  过滤历史重复：{len(removed_dup)} 条")
    if removed_dup:
        log(f"    {removed_dup[:3]}")
    log(f"  剩余候选：{len(passed)} 条")

    return passed


# ========== 图片广告过滤 ==========
def _is_ad_image(url: str) -> bool:
    """
    判断图片 URL 是否为广告/无关图片，返回 True 表示应被过滤掉。
    过滤规则：
    1. URL 含广告关键词
    2. 来自已知广告/追踪域名
    3. 疑似头像/logo/icon（小图标类）
    4. 含追踪参数
    5. 社交平台头像或用户图片
    """
    if not url:
        return True

    url_lower = url.lower()

    # ── 规则1：URL 路径含广告关键词 ──
    AD_PATH_KEYWORDS = [
        "/ad/", "/ads/", "/adv/", "/advert/", "/advertisement/",
        "/banner/", "/banners/", "/sponsor/", "/sponsored/",
        "/promo/", "/promotion/", "/track/", "/tracking/",
        "/click/", "/clickthrough/", "/redirect/",
        "/logo/", "/logos/", "/icon/", "/icons/",
        "/avatar/", "/avatars/", "/headimg/", "/profile/",
        "/thumb/", "/thumbnail/", "/emoji/",
        "favicon", "watermark",
    ]
    for kw in AD_PATH_KEYWORDS:
        if kw in url_lower:
            return True

    # ── 规则2：已知广告/追踪域名 ──
    AD_DOMAINS = [
        "doubleclick.net", "googlesyndication.com", "googleadservices.com",
        "adservice.google", "amazon-adsystem.com", "facebook.com/tr",
        "connect.facebook.net", "pixel.twitter.com", "analytics.twitter.com",
        "bat.bing.com", "adnxs.com", "taobao.com", "jd.com/adv",
        "pingan.com.cn", "cpro.baidustatic.com", "pos.baidu.com",
        "tianqi.com", "qzone.qq.com", "thirdqq.qpic.cn",
        "wx.qlogo.cn", "thirdwx.qlogo.cn",   # 微信头像
        "q.qlogo.cn",                          # QQ 头像
    ]
    for domain in AD_DOMAINS:
        if domain in url_lower:
            return True

    # ── 规则3：URL 含追踪参数特征 ──
    AD_PARAM_KEYWORDS = [
        "utm_source", "utm_medium", "utm_campaign",
        "adid=", "adunit=", "adtype=", "adslot=",
        "clickid=", "affilate=",
    ]
    for kw in AD_PARAM_KEYWORDS:
        if kw in url_lower:
            return True

    # ── 规则4：图片文件名含广告/无关特征 ──
    import os as _os
    fname = _os.path.basename(url_lower.split("?")[0])
    AD_FNAME_KEYWORDS = [
        "ad_", "_ad.", "adv_", "banner_", "logo_", "icon_",
        "avatar_", "head_", "qrcode", "qr_code", "erweima",
        "weixin", "wechat_", "app_download", "appstore",
    ]
    for kw in AD_FNAME_KEYWORDS:
        if kw in fname:
            return True

    # ── 规则5：图片后缀为 gif（多为 banner 动图或追踪像素）──
    if fname.endswith(".gif"):
        return True

    return False


def _filter_images(raw_urls: list, max_count: int = 3) -> list:
    """
    过滤广告图，返回最多 max_count 张干净图片。
    """
    result = []
    filtered_out = []
    for url in raw_urls:
        if _is_ad_image(url):
            filtered_out.append(url)
        else:
            result.append(url)
        if len(result) >= max_count:
            break
    if filtered_out:
        log(f"  🚫 过滤广告/无关图片 {len(filtered_out)} 张: {[u[:50] for u in filtered_out]}")
    return result


# ========== Step 3: 搜索原文内容 + 图片 ==========
def search_news_detail(title):
    """
    通过 Tavily 搜索原文详情 + 图片
    返回 (content_text, image_urls)
    """
    content_text = ""
    image_urls = []

    if not TAVILY_API_KEY:
        log(f"  ⚠️  未配置 TAVILY_API_KEY")
        return content_text, image_urls

    try:
        payload = json.dumps({
            "api_key": TAVILY_API_KEY,
            "query": title,
            "search_depth": "advanced",
            "max_results": 5,
            "include_answer": True,
            "include_raw_content": False,
            "include_images": True,
        }).encode("utf-8")

        req = urllib.request.Request(
            "https://api.tavily.com/search",
            data=payload,
            headers={"Content-Type": "application/json"},
            method="POST"
        )
        with urllib.request.urlopen(req, timeout=20) as resp:
            data = json.loads(resp.read().decode("utf-8"))

        parts = []
        if data.get("answer"):
            parts.append(f"【核心答案】{data['answer']}")

        for result in data.get("results", [])[:5]:
            snippet = result.get("content", "")[:400]
            url = result.get("url", "")
            title_r = result.get("title", "")
            if snippet:
                parts.append(f"【来源：{title_r}】\n{snippet}")

        content_text = "\n\n".join(parts)[:2000]

        # 提取图片（多抓一些，过滤后保留干净的）
        raw_imgs = []
        for img in data.get("images", [])[:10]:
            if isinstance(img, str) and img.startswith("http"):
                raw_imgs.append(img)
            elif isinstance(img, dict):
                u = img.get("url", "")
                if u.startswith("http"):
                    raw_imgs.append(u)

        image_urls = _filter_images(raw_imgs, max_count=3)
        log(f"  ✅ 搜索完成：{len(parts)} 段内容，图片 {len(raw_imgs)} 张 → 过滤后 {len(image_urls)} 张")

    except Exception as e:
        log(f"  搜索失败（{title[:20]}）: {e}")

    return content_text, image_urls


def search_images_fallback(title, keywords=""):
    """
    备用图片搜索：如果主搜索没找到图片，使用多个来源搜索
    keywords: 额外的关键词，如"科技"、"AI"等
    """
    log(f"  🔍 备用图片搜索：{title[:30]}...")
    
    search_query = f"{title} {keywords}".strip()
    image_urls = []
    
    # 方法1：使用Tavily搜索图片（focus on images）
    if TAVILY_API_KEY:
        try:
            payload = json.dumps({
                "api_key": TAVILY_API_KEY,
                "query": search_query,
                "search_depth": "basic",
                "max_results": 3,
                "include_images": True,
            }).encode("utf-8")
            
            req = urllib.request.Request(
                "https://api.tavily.com/search",
                data=payload,
                headers={"Content-Type": "application/json"},
                method="POST"
            )
            with urllib.request.urlopen(req, timeout=15) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            
            raw_imgs = []
            for img in data.get("images", [])[:15]:
                if isinstance(img, str) and img.startswith("http"):
                    raw_imgs.append(img)
                elif isinstance(img, dict):
                    u = img.get("url", "")
                    if u.startswith("http"):
                        raw_imgs.append(u)
            
            image_urls = _filter_images(raw_imgs, max_count=3)
            if image_urls:
                log(f"  ✅ 备用搜索找到 {len(image_urls)} 张图片")
                return image_urls
        except Exception as e:
            log(f"  ⚠️  备用图片搜索失败: {str(e)[:50]}")
    
    # 方法2：使用通用关键词搜索
    generic_queries = [
        f"{title[:20]} 新闻图片",
        f"{title[:20]} 配图",
        title[:15],
    ]
    
    for query in generic_queries:
        try:
            if not TAVILY_API_KEY:
                break
            
            payload = json.dumps({
                "api_key": TAVILY_API_KEY,
                "query": query,
                "search_depth": "basic",
                "max_results": 2,
                "include_images": True,
            }).encode("utf-8")
            
            req = urllib.request.Request(
                "https://api.tavily.com/search",
                data=payload,
                headers={"Content-Type": "application/json"},
                method="POST"
            )
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            
            raw_imgs = []
            for img in data.get("images", [])[:10]:
                if isinstance(img, str) and img.startswith("http"):
                    raw_imgs.append(img)
                elif isinstance(img, dict):
                    u = img.get("url", "")
                    if u.startswith("http"):
                        raw_imgs.append(u)
            
            image_urls = _filter_images(raw_imgs, max_count=3)
            if image_urls:
                log(f"  ✅ 通用搜索找到 {len(image_urls)} 张图片（查询：{query}）")
                return image_urls
        except Exception as e:
            continue
    
    log(f"  ❌ 所有方法都未找到可用图片")
    return []


# ========== Step 4: AI 改写标题 + 生成图文内容 ==========
def extract_images_from_markdown(md_content):
    """
    从 Markdown 中提取图片 URL（![描述](url) 格式）
    """
    import re
    urls = []
    # 匹配 ![...](http://... 或 https://...)
    pattern = r'!\[([^\]]*)\]\((https?://[^)]+)\)'
    matches = re.findall(pattern, md_content)
    for desc, url in matches:
        if url not in urls:
            urls.append(url)
    return urls[:5]  # 最多提取5张


def ai_rewrite(original_title, search_content, image_urls):
    """
    调用 AI 改写标题并生成图文内容
    返回 (new_title, content_text, md_content)
    """
    # 图片占位符说明
    img_hint = ""
    if image_urls:
        img_hint = f"\n\n【可用图片URL（在合适位置插入）】\n" + "\n".join(
            [f"图片{i+1}: {url}" for i, url in enumerate(image_urls[:3])]
        )

    prompt = f"""你是一位今日头条爆款文章写手，请根据以下信息完成任务：

【原始热榜标题】
{original_title}

【相关背景信息（来自网络搜索）】
{search_content if search_content else "（无额外背景信息，请基于标题内容发挥）"}
{img_hint}

【任务要求】

1. 改写标题：
   - 在原标题基础上稍作改动，不能与原标题完全相同
   - 适当增加emoji表情（1~2个）
   - 总字数不超过30字
   - 风格：有吸引力、制造悬念或强调重要性

2. 生成正文（图文并茂）：
   - 字数：500~700字
   - 风格：头条爆款，逻辑清晰，语言通俗有力，有观点有态度
   - 结构：
     * 第1段：开门见山，1~2句话抓住读者（最重要的事实或最强观点）
     * 第2~4段：核心内容展开，每段有小标题（用**加粗**），每段配1个emoji
     * 第5段：结尾点评或引导读者评论互动
   - 如果有可用图片URL，在第1段后插入第一张图片，格式：![图片描述](图片URL)
   - 如果有多张图片，可在后续段落适当插入

请严格按照以下JSON格式输出，不要输出任何其他内容：
{{
  "title": "改写后的标题（含emoji，不超过30字）",
  "content": "纯文本正文内容（换行用\\n，不含图片标记）",
  "md_content": "完整Markdown格式正文（含图片插入，换行用\\n）"
}}"""

    try:
        payload = json.dumps({
            "model": MODEL_ID,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 2000,
            "temperature": 0.8,
        }).encode("utf-8")

        req = urllib.request.Request(
            f"{AI_API_URL}/chat/completions",
            data=payload,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {AI_API_KEY}",
            },
            method="POST"
        )

        with urllib.request.urlopen(req, timeout=90) as resp:
            data = json.loads(resp.read().decode("utf-8"))

        text = data["choices"][0]["message"]["content"].strip()

        # 去掉 markdown code block 包装
        text = re.sub(r'^```json\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r'^```\s*$', '', text, flags=re.MULTILINE)
        text = text.strip()

        result = json.loads(text)
        new_title  = result.get("title", "").strip()
        content    = result.get("content", "").strip()
        md_content = result.get("md_content", content).strip()

        # 标题字数检查
        if len(new_title) > 35:
            new_title = new_title[:32] + "..."

        return new_title, content, md_content

    except json.JSONDecodeError as e:
        log(f"  AI JSON解析失败: {e}")
        log(f"  原始输出片段: {text[:300] if 'text' in dir() else 'N/A'}")
        return "", "", ""
    except Exception as e:
        log(f"  AI调用失败: {e}")
        return "", "", ""


def ai_review_and_fix(title, content, md_content):
    """
    AI质量审核：检查并修复格式、行文、排版问题
    返回 (fixed_title, fixed_content, fixed_md_content, issues_found)
    """
    prompt = f"""你是一位专业的内容审核编辑，请审核以下今日头条文章，找出问题并修复。

【文章标题】
{title}

【文章正文（Markdown格式）】
{md_content}

【审核要点】
1. 格式检查：
   - 标题是否有乱码、多余符号、重复emoji
   - 段落格式是否正确（小标题需要**加粗**）
   - emoji使用是否合理（每段最多1个，避免过度使用）
   - 是否有未闭合的标记符号
   - **必须保留所有图片标记**：![图片描述](图片路径)

2. 内容质量：
   - 是否有乱码或错乱文字
   - 语句是否流畅通顺
   - 逻辑是否连贯
   - 是否有重复内容
   - 段落长度是否合理（每段2-4句话）

3. 头条号规范：
   - 避免敏感词汇
   - 标题不超过30字
   - 正文500-700字
   - 结尾需要互动引导

【重要】修复时必须保留原文中的所有图片标记！格式：![描述](路径)

请输出JSON格式（不要输出其他内容）：
{{
  "has_issues": true/false,
  "issues": ["问题1描述", "问题2描述"],
  "fixed_title": "修复后的标题",
  "fixed_content": "修复后的纯文本正文",
  "fixed_md_content": "修复后的Markdown正文"
}}"""

    try:
        payload = json.dumps({
            "model": MODEL_ID,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 2500,
            "temperature": 0.3,  # 降低温度，让输出更稳定
        }).encode("utf-8")

        req = urllib.request.Request(
            f"{AI_API_URL}/chat/completions",
            data=payload,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {AI_API_KEY}",
            },
            method="POST"
        )

        with urllib.request.urlopen(req, timeout=90) as resp:
            data = json.loads(resp.read().decode("utf-8"))

        text = data["choices"][0]["message"]["content"].strip()
        
        # 去掉 markdown code block 包装
        text = re.sub(r'^```json\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r'^```\s*$', '', text, flags=re.MULTILINE)
        text = text.strip()

        result = json.loads(text)
        
        has_issues = result.get("has_issues", False)
        issues = result.get("issues", [])
        fixed_title = result.get("fixed_title", title).strip()
        fixed_content = result.get("fixed_content", content).strip()
        fixed_md_content = result.get("fixed_md_content", md_content).strip()
        
        # 标题字数再次检查
        if len(fixed_title) > 35:
            fixed_title = fixed_title[:32] + "..."
        
        return fixed_title, fixed_content, fixed_md_content, issues
        
    except Exception as e:
        log(f"  ⚠️  AI审核失败，使用原内容: {e}")
        return title, content, md_content, []


# ========== Step 5: 保存 JSON + Markdown ==========
def save_article(index, original_title, new_title, content, md_content, image_urls, search_content):
    """保存文章为 JSON（发布用）和 MD（存档用）"""
    today = datetime.now().strftime("%Y%m%d")

    # 清理标题中不能用于文件名的字符
    safe_title = re.sub(r'[\\/:*?"<>|\n\r]', '_', new_title)
    safe_title = re.sub(r'[\U0001F000-\U0001FFFF]', '', safe_title).strip()  # 去除emoji
    safe_title = safe_title[:40]
    
    # 下载图片到本地
    local_image_paths = []
    if image_urls:
        for i, img_url in enumerate(image_urls):
            try:
                import urllib.request as ur
                req = ur.Request(img_url, headers={
                    'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36',
                    'Referer': 'https://www.toutiao.com/'
                })
                with ur.urlopen(req, timeout=10) as resp:
                    img_data = resp.read()
                
                # 保存到MD_DIR
                img_filename = f"{today}_{index:02d}_img{i+1}.jpg"
                img_path = os.path.join(MD_DIR, img_filename)
                with open(img_path, 'wb') as f:
                    f.write(img_data)
                
                local_image_paths.append(img_path)
                
                # 替换MD中的图片链接为本地路径
                md_content = md_content.replace(img_url, img_path)
                log(f"  📥 图片{i+1}下载成功: {len(img_data)} bytes")
                
            except Exception as e:
                log(f"  ⚠️  图片{i+1}下载失败，跳过: {str(e)[:50]}")
                # 不保留外链URL，只使用成功下载的本地图片
    
    # 只使用成功下载的本地图片
    if local_image_paths:
        image_urls = local_image_paths
        log(f"  ✅ 成功下载 {len(local_image_paths)}/{len(image_urls) if image_urls else 0} 张图片")
    else:
        image_urls = []  # 没有可用图片
        log(f"  ⚠️  所有图片下载失败，文章将无图片")

    # 保存 JSON
    json_path = os.path.join(OUTPUT_DIR, f"{today}_{index:02d}.json")
    article = {
        "index": index,
        "date": today,
        "original_title": original_title,
        "title": new_title,
        "content": content,
        "md_content": md_content,
        "image_urls": image_urls,
        "created_at": datetime.now().isoformat(),
        "published": False,
    }
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(article, f, ensure_ascii=False, indent=2)

    # 保存 Markdown（图文存档）
    md_path = os.path.join(MD_DIR, f"{today}_{index:02d}_{safe_title}.md")
    md_full = f"""# {new_title}

> 原始热榜：{original_title}
> 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}

---

{md_content}

---

<details>
<summary>搜索原始资料</summary>

{search_content[:500] if search_content else '无'}

</details>
"""
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_full)

    log(f"  💾 JSON: {os.path.basename(json_path)}")
    log(f"  📄 MD:   {os.path.basename(md_path)}")
    
    # 返回修改后的md_content（包含本地图片路径）和image_urls（本地路径）
    return json_path, md_path, md_content, image_urls


def save_article_as_word(index, title, content, md_content, image_urls):
    """将文章保存为Word文档格式（带图片）"""
    today = datetime.now().strftime("%Y%m%d")
    
    # 清理标题中不能用于文件名的字符
    safe_title = re.sub(r'[\\/:*?"<>|\n\r]', '_', title)
    safe_title = re.sub(r'[\U0001F000-\U0001FFFF]', '', safe_title).strip()
    safe_title = safe_title[:40]
    
    docx_path = os.path.join(DOCX_DIR, f"{today}_{index:02d}_{safe_title}.docx")
    
    try:
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 解析md_content并添加到Word
        lines = md_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # 跳过空行
            if not line:
                i += 1
                continue
            
            # 图片
            if line.startswith('!['):
                img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
                if img_match:
                    img_path = img_match.group(2)
                    if os.path.exists(img_path):
                        try:
                            # 添加图片（宽度6英寸，保持比例）
                            doc.add_picture(img_path, width=Inches(6))
                            # 图片居中
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            log(f"  ⚠️  Word图片插入失败: {e}")
            
            # 小标题（**加粗**）
            elif line.startswith('**') and line.endswith('**'):
                heading_text = line.strip('*').strip()
                para = doc.add_paragraph()
                run = para.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # 普通段落
            else:
                # 处理行内格式
                para = doc.add_paragraph()
                # 简单处理：移除**标记，保留文字
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                para.add_run(text)
                para.paragraph_format.line_spacing = 1.5
            
            i += 1
        
        # 保存文档
        doc.save(docx_path)
        log(f"  📄 Word: {os.path.basename(docx_path)}")
        return docx_path
        
    except Exception as e:
        log(f"  ⚠️  Word文档生成失败: {e}")
        return None


# ========== 主流程 ==========
def main():
    log("=" * 55)
    log("🚀 头条热榜新闻自动生成脚本 v2.1 启动")
    log("=" * 55)

    # 清理7日以上的旧文档
    cleanup_old_documents(days=7)
    log("")

    # 加载历史标题
    history_titles = load_history_titles()
    log(f"📚 已加载历史标题 {len(history_titles)} 条（用于去重）")

    # Step 1: 抓取热榜（多抓一些备用）
    raw_titles, raw_items = fetch_toutiao_hot(max_count=FETCH_POOL)
    if not raw_titles:
        log("❌ 未能获取热榜数据，退出")
        sys.exit(1)

    log(f"\n📋 原始热榜（共 {len(raw_titles)} 条）：")
    for i, t in enumerate(raw_titles, 1):
        print(f"  {i:2d}. {t}")

    # Step 2: 过滤
    log("\n🔍 开始过滤...")
    candidates = filter_titles(raw_titles, history_titles)

    log(f"\n✅ 候选新闻（共 {len(candidates)} 条）：")
    for i, t in enumerate(candidates, 1):
        print(f"  {i:2d}. {t}")

    if not candidates:
        log("❌ 过滤后无候选标题，退出")
        sys.exit(0)

    # 确保候选数量足够达到目标
    if len(candidates) < TARGET_MIN:
        log(f"⚠️  候选数量 {len(candidates)} < 目标最小值 {TARGET_MIN}，将处理所有候选")
    work_titles = candidates[:TARGET_MAX]  # 最多处理 TARGET_MAX 条

    log(f"\n🎯 本次将处理 {len(work_titles)} 条（目标 {TARGET_MIN}~{TARGET_MAX} 篇）")

    # Step 3+4+5: 逐条搜索 + AI 改写 + 保存
    articles = []
    article_index = 1

    for i, title in enumerate(work_titles, 1):
        log(f"\n{'─'*45}")
        log(f"[{i}/{len(work_titles)}] 处理：{title}")

        # 搜索原文内容 + 图片
        log(f"  🔍 搜索原文内容和图片...")
        search_content, image_urls = search_news_detail(title)
        
        # 如果没有图片，尝试备用搜索
        if not image_urls or len(image_urls) < MIN_IMAGES:
            log(f"  ⚠️  图片不足（{len(image_urls)}/{MIN_IMAGES}），尝试备用搜索...")
            fallback_images = search_images_fallback(title, keywords="新闻 热点")
            if fallback_images:
                image_urls.extend(fallback_images)
                image_urls = image_urls[:3]  # 最多保留3张
                log(f"  ✅ 补充后共 {len(image_urls)} 张图片")
        
        # 检查图片数量是否满足要求
        if not image_urls or len(image_urls) < MIN_IMAGES:
            log(f"  ❌ 图片不足（{len(image_urls)}/{MIN_IMAGES}），跳过此条")
            log(f"     原因：每篇文章至少需要 {MIN_IMAGES} 张图片")
            continue

        # AI 改写
        log(f"  🤖 AI 改写（图文并茂）...")
        new_title, content, md_content = ai_rewrite(title, search_content, image_urls)

        if not new_title or not content:
            log(f"  ❌ AI 改写失败，跳过此条")
            continue

        # AI 质量审核
        log(f"  🔍 AI 质量审核...")
        fixed_title, fixed_content, fixed_md_content, issues = ai_review_and_fix(
            new_title, content, md_content
        )
        
        # 验证图片是否被保留
        original_imgs = re.findall(r'!\[[^\]]*\]\([^)]+\)', md_content)
        fixed_imgs = re.findall(r'!\[[^\]]*\]\([^)]+\)', fixed_md_content)
        
        if len(original_imgs) > len(fixed_imgs):
            log(f"  ⚠️  审核时{len(original_imgs) - len(fixed_imgs)}张图片被删除，恢复中...")
            
            # 找出丢失的图片
            fixed_img_set = set(fixed_imgs)
            missing_imgs = [img for img in original_imgs if img not in fixed_img_set]
            
            if missing_imgs:
                # 策略：将丢失的图片均匀插入到段落之间
                lines = fixed_md_content.split('\n')
                
                # 找到所有段落位置（以**开头的小标题后）
                paragraph_positions = []
                for i, line in enumerate(lines):
                    if line.strip().startswith('**') and i + 1 < len(lines):
                        paragraph_positions.append(i + 1)
                
                # 如果没有段落标题，就在开头插入
                if not paragraph_positions:
                    for i, line in enumerate(lines):
                        if line.strip() and not line.startswith('#') and not line.startswith('>'):
                            paragraph_positions.append(i + 1)
                            break
                
                # 将图片均匀分布
                if paragraph_positions:
                    img_interval = max(1, len(paragraph_positions) // len(missing_imgs))
                    insert_positions = paragraph_positions[::img_interval][:len(missing_imgs)]
                    
                    # 从后往前插入（避免位置偏移）
                    for pos, img in zip(reversed(insert_positions), reversed(missing_imgs)):
                        lines.insert(pos, '')
                        lines.insert(pos + 1, img)
                    
                    fixed_md_content = '\n'.join(lines)
                    log(f"  ✅ 已恢复 {len(missing_imgs)} 张图片")
        
        if issues:
            log(f"  📝 发现 {len(issues)} 个问题并已修复：")
            for issue in issues[:3]:  # 最多显示3个问题
                log(f"     - {issue}")
            new_title = fixed_title
            content = fixed_content
            md_content = fixed_md_content
        else:
            log(f"  ✅ 内容质量检查通过")

        # 从 Markdown 中提取图片（如果搜索结果中没有）
        if not image_urls:
            extracted_urls = extract_images_from_markdown(md_content)
            if extracted_urls:
                image_urls = extracted_urls
                log(f"  📸 从 Markdown 提取到 {len(image_urls)} 张图片")

        log(f"  ✅ 处理完成")
        log(f"     新标题：{new_title}")
        log(f"     字数：{len(content)} 字 | 图片：{len(image_urls)} 张")

        # 保存文件（图片下载并替换链接）
        json_path, md_path, md_content_with_local_images, local_image_urls = save_article(
            article_index, title, new_title, content, md_content,
            image_urls, search_content
        )
        
        # 生成Word文档（使用本地图片路径）
        docx_path = save_article_as_word(
            article_index, new_title, content, md_content_with_local_images, local_image_urls
        )

        # 记录历史（用新标题去重，防止明天重复）
        save_history_title(title)
        save_history_title(new_title)

        articles.append({
            "index": article_index,
            "original_title": title,
            "title": new_title,
            "content": content,
            "md_content": md_content,
            "image_urls": image_urls,
            "json_path": json_path,
            "md_path": md_path,
        })
        article_index += 1

        # 达到目标最大值时停止
        if len(articles) >= TARGET_MAX:
            log(f"\n✅ 已达到目标最大值 {TARGET_MAX} 篇，停止处理")
            break

        # 请求间隔
        if i < len(work_titles):
            time.sleep(2)

    # ========== 汇总 ==========
    log("\n" + "=" * 55)
    log(f"🎉 处理完成！共生成 {len(articles)} 篇文章")
    log("=" * 55)

    for a in articles:
        print(f"\n{'━'*45}")
        print(f"📰 原始：{a['original_title']}")
        print(f"✏️  新标：{a['title']}")
        print(f"📊 字数：{len(a['content'])} 字 | 图片：{len(a['image_urls'])} 张")
        print(f"📄 内容预览：{a['content'][:120]}...")
        print(f"💾 MD文件：{os.path.basename(a['md_path'])}")

    if len(articles) < TARGET_MIN:
        log(f"\n⚠️  生成数量 {len(articles)} < 目标最小值 {TARGET_MIN}，建议检查过滤规则")
    else:
        log(f"\n✅ 目标达成：{len(articles)} 篇（目标 {TARGET_MIN}~{TARGET_MAX}）")

    log(f"\n📁 JSON文章：{OUTPUT_DIR}")
    log(f"📁 MD存档：{MD_DIR}")
    log(f"📋 历史记录：{HISTORY_FILE}")

    return articles


if __name__ == "__main__":
    main()
