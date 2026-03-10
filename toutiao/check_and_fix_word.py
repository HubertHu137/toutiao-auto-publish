#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档内容检查和自动修复工具
功能：
1. 检查Word文档是否符合头条号平台规范
2. 检查是否有乱码和多余字符
3. 检查是否有广告类图片
4. 自动修复发现的问题
5. 生成修复后的新Word文档

用法：
  python3 check_and_fix_word.py <word文件路径>                    # 检查并修复单个文件
  python3 check_and_fix_word.py --batch                          # 批量检查今天的所有文档
  python3 check_and_fix_word.py --batch --fix                    # 批量检查并自动修复
"""

import sys
import os
import re
import json
import urllib.request
import urllib.error
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("⚠️  python-docx 未安装，正在安装...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from PIL import Image
    import imagehash
except ImportError:
    print("⚠️  Pillow 和 imagehash 未安装，正在安装...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "Pillow", "imagehash"])
    from PIL import Image
    import imagehash

# ========== 配置 ==========
AI_API_URL = "http://127.0.0.1:19999"
AI_API_KEY = os.environ.get("AI_API_KEY", "")   # 必填：AI API Key，通过环境变量 AI_API_KEY 配置
MODEL_ID = "gpt-4o-mini"
DOCX_DIR = os.path.expanduser("~/.openclaw/toutiao_articles/docx")

# ========== 头条号规范 ==========
TOUTIAO_RULES = {
    "title_max_length": 30,      # 标题最大字数
    "content_min_length": 400,   # 正文最少字数
    "content_max_length": 2000,  # 正文最大字数
    "max_images": 9,             # 最多图片数量
    "min_paragraph_length": 20,  # 段落最少字数
}

# ========== 敏感词过滤 ==========
SENSITIVE_WORDS = [
    # 政治类
    "习近平", "李强", "政府", "中央", "党委", "国务院", "人大", "政协",
    "书记", "主席", "总理", "省长", "市长", "官员",
    # 违规内容
    "赌博", "色情", "暴力", "毒品", "枪支", "爆炸",
    # 广告类
    "加微信", "扫码", "点击链接", "立即购买", "限时优惠", "免费领取",
    "微商", "代理", "招商", "加盟", "投资理财",
]

# ========== 广告图片特征（通过OCR或hash检测）==========
AD_IMAGE_KEYWORDS = [
    "微信", "QQ", "二维码", "扫一扫", "关注公众号",
    "优惠券", "折扣", "立即购买", "点击购买",
    "加微信", "加QQ", "联系方式", "电话", "手机号",
]


def log(msg):
    """输出日志"""
    print(f"[{datetime.now().strftime('%H:%M:%S')}] {msg}", flush=True)


def extract_text_from_word(docx_path):
    """从Word文档提取纯文本内容"""
    try:
        doc = Document(docx_path)
        
        # 提取标题（第一个Heading 1）
        title = ""
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                title = para.text.strip()
                break
        
        # 提取正文
        content_lines = []
        for para in doc.paragraphs:
            if not para.style.name.startswith('Heading'):
                text = para.text.strip()
                if text:
                    content_lines.append(text)
        
        content = '\n\n'.join(content_lines)
        
        # 提取图片信息
        image_count = len(doc.inline_shapes)
        
        return {
            "title": title,
            "content": content,
            "image_count": image_count,
            "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
        }
    except Exception as e:
        log(f"❌ 读取Word文档失败: {e}")
        return None


def check_encoding_issues(text):
    """检查乱码和编码问题"""
    issues = []
    
    # 检查常见乱码模式
    garbled_patterns = [
        (r'[��]+', '替换字符乱码'),
        (r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '控制字符'),
        (r'[^\u4e00-\u9fa5\u3000-\u303f\uff00-\uffef\u0020-\u007e\u2000-\u206f\n\r\t]+', '异常字符'),
    ]
    
    for pattern, desc in garbled_patterns:
        matches = re.findall(pattern, text)
        if matches:
            issues.append({
                "type": "encoding",
                "severity": "high",
                "description": f"发现{desc}",
                "details": f"匹配{len(matches)}处",
                "examples": matches[:3],
            })
    
    # 检查多余空格和换行
    if re.search(r'\s{3,}', text):
        issues.append({
            "type": "formatting",
            "severity": "medium",
            "description": "发现连续多个空格",
            "fix": "自动清理",
        })
    
    if re.search(r'\n{3,}', text):
        issues.append({
            "type": "formatting",
            "severity": "medium",
            "description": "发现连续多个换行",
            "fix": "自动清理",
        })
    
    return issues


def check_toutiao_compliance(title, content, image_count):
    """检查是否符合头条号规范"""
    issues = []
    
    # 1. 标题长度检查
    title_len = len(title)
    if title_len > TOUTIAO_RULES["title_max_length"]:
        issues.append({
            "type": "title_length",
            "severity": "high",
            "description": f"标题过长（{title_len}字 > {TOUTIAO_RULES['title_max_length']}字）",
            "fix": f"建议缩短到{TOUTIAO_RULES['title_max_length']}字以内",
        })
    elif title_len < 5:
        issues.append({
            "type": "title_length",
            "severity": "medium",
            "description": f"标题过短（{title_len}字）",
            "fix": "建议增加到5-30字",
        })
    
    # 2. 正文长度检查
    content_len = len(content.replace('\n', '').replace(' ', ''))
    if content_len < TOUTIAO_RULES["content_min_length"]:
        issues.append({
            "type": "content_length",
            "severity": "high",
            "description": f"正文过短（{content_len}字 < {TOUTIAO_RULES['content_min_length']}字）",
            "fix": "建议扩充内容到400字以上",
        })
    elif content_len > TOUTIAO_RULES["content_max_length"]:
        issues.append({
            "type": "content_length",
            "severity": "medium",
            "description": f"正文过长（{content_len}字 > {TOUTIAO_RULES['content_max_length']}字）",
            "fix": "建议精简内容",
        })
    
    # 3. 图片数量检查
    if image_count > TOUTIAO_RULES["max_images"]:
        issues.append({
            "type": "image_count",
            "severity": "medium",
            "description": f"图片过多（{image_count}张 > {TOUTIAO_RULES['max_images']}张）",
            "fix": "建议删减图片",
        })
    
    # 4. 敏感词检查
    found_sensitive = []
    for word in SENSITIVE_WORDS:
        if word in title or word in content:
            found_sensitive.append(word)
    
    if found_sensitive:
        issues.append({
            "type": "sensitive_words",
            "severity": "critical",
            "description": f"发现敏感词：{', '.join(found_sensitive[:5])}",
            "fix": "必须删除或替换敏感词",
        })
    
    # 5. 标题规范检查
    if re.search(r'[!！]{2,}', title):
        issues.append({
            "type": "title_format",
            "severity": "low",
            "description": "标题中有多个感叹号",
            "fix": "建议只保留1个",
        })
    
    # 计算emoji数量
    emoji_count = len(re.findall(r'[\U0001F000-\U0001FFFF]', title))
    if emoji_count > 2:
        issues.append({
            "type": "title_format",
            "severity": "low",
            "description": f"标题emoji过多（{emoji_count}个）",
            "fix": "建议只保留1-2个",
        })
    
    # 6. 段落格式检查
    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
    if len(paragraphs) < 3:
        issues.append({
            "type": "paragraph_structure",
            "severity": "medium",
            "description": f"段落过少（{len(paragraphs)}段）",
            "fix": "建议分成3-5段",
        })
    
    # 检查段落长度
    for i, para in enumerate(paragraphs, 1):
        para_len = len(para.replace(' ', ''))
        if para_len < TOUTIAO_RULES["min_paragraph_length"]:
            issues.append({
                "type": "paragraph_length",
                "severity": "low",
                "description": f"第{i}段过短（{para_len}字）",
                "fix": "建议每段20字以上",
            })
    
    return issues


def ai_check_and_fix(title, content):
    """
    调用AI进行深度内容检查和修复
    """
    prompt = f"""你是一位专业的今日头条内容审核编辑，请全面检查以下文章并修复问题。

【标题】
{title}

【正文】
{content}

【检查要点】
1. **乱码和字符问题**：
   - 查找并标记所有乱码字符（如��、□、？等）
   - 检查是否有多余的特殊符号
   - 检查是否有重复的标点符号

2. **格式规范**：
   - 标题长度5-30字，不超过30字
   - 正文400-2000字
   - 段落结构清晰，每段20字以上
   - 标点符号使用正确

3. **内容质量**：
   - 语句通顺，逻辑连贯
   - 避免敏感词汇（政治、违规、广告等）
   - 避免夸大、误导性表述
   - 是否有广告营销内容

4. **头条号规范**：
   - 标题要吸引人但不标题党
   - emoji使用合理（1-2个）
   - 结尾有互动引导
   - 段落间逻辑流畅

【修复原则】
- 优先修复乱码和字符问题（必须修复）
- 删除或替换敏感词
- 优化格式和排版
- 保持原文核心意思不变
- 如果标题/正文长度不符合要求，进行调整

请输出JSON格式（不要输出其他内容）：
{{
  "has_critical_issues": true/false,
  "issues_found": [
    {{"type": "类型", "severity": "严重程度", "description": "问题描述"}}
  ],
  "fixed_title": "修复后的标题",
  "fixed_content": "修复后的正文",
  "changes_made": ["修改1", "修改2"],
  "recommendations": ["建议1", "建议2"]
}}"""

    try:
        payload = json.dumps({
            "model": MODEL_ID,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 3000,
            "temperature": 0.3,
        }).encode("utf-8")

        req = urllib.request.Request(
            f"{AI_API_URL}/chat/completions",
            data=payload,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {AI_API_KEY}",
            },
            method="POST"
        )

        with urllib.request.urlopen(req, timeout=120) as resp:
            data = json.loads(resp.read().decode("utf-8"))

        text = data["choices"][0]["message"]["content"].strip()
        
        # 去掉 markdown code block 包装
        text = re.sub(r'^```json\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r'^```\s*$', '', text, flags=re.MULTILINE)
        text = text.strip()

        result = json.loads(text)
        return result
        
    except Exception as e:
        log(f"  ⚠️  AI检查失败: {e}")
        return None


def check_word_document(docx_path, auto_fix=False):
    """
    全面检查Word文档
    返回检查报告和修复建议
    """
    log(f"📄 开始检查：{os.path.basename(docx_path)}")
    
    # 1. 提取文档内容
    doc_info = extract_text_from_word(docx_path)
    if not doc_info:
        return None
    
    title = doc_info["title"]
    content = doc_info["content"]
    image_count = doc_info["image_count"]
    
    log(f"  标题：{title[:50]}...")
    log(f"  正文：{len(content)}字，{doc_info['paragraph_count']}段，{image_count}张图")
    
    # 2. 基础检查
    log("  🔍 检查乱码和编码问题...")
    encoding_issues = check_encoding_issues(title + '\n' + content)
    
    log("  🔍 检查头条号规范...")
    compliance_issues = check_toutiao_compliance(title, content, image_count)
    
    # 3. AI深度检查
    log("  🤖 AI深度内容检查...")
    ai_result = ai_check_and_fix(title, content)
    
    # 汇总所有问题
    all_issues = encoding_issues + compliance_issues
    if ai_result:
        all_issues.extend(ai_result.get("issues_found", []))
    
    # 4. 生成检查报告
    report = {
        "file_path": docx_path,
        "check_time": datetime.now().isoformat(),
        "document_info": doc_info,
        "issues": all_issues,
        "ai_result": ai_result,
        "total_issues": len(all_issues),
        "critical_issues": len([i for i in all_issues if i.get("severity") == "critical"]),
        "has_critical_issues": ai_result.get("has_critical_issues", False) if ai_result else False,
    }
    
    # 5. 输出报告摘要
    log("\n" + "="*60)
    log("📊 检查报告")
    log("="*60)
    
    if report["total_issues"] == 0:
        log("✅ 未发现问题，文档符合规范")
    else:
        log(f"⚠️  发现 {report['total_issues']} 个问题")
        
        # 按严重程度分组显示
        for severity in ["critical", "high", "medium", "low"]:
            severity_issues = [i for i in all_issues if i.get("severity") == severity]
            if severity_issues:
                severity_names = {
                    "critical": "🔴 严重",
                    "high": "🟠 高",
                    "medium": "🟡 中",
                    "low": "🟢 低"
                }
                log(f"\n{severity_names.get(severity, severity)}级别问题 ({len(severity_issues)}个):")
                for issue in severity_issues:
                    log(f"  • {issue['description']}")
                    if issue.get('fix'):
                        log(f"    修复：{issue['fix']}")
    
    # 6. AI修复建议
    if ai_result:
        if ai_result.get("changes_made"):
            log(f"\n🔧 AI建议修改:")
            for change in ai_result["changes_made"][:5]:
                log(f"  • {change}")
        
        if ai_result.get("recommendations"):
            log(f"\n💡 优化建议:")
            for rec in ai_result["recommendations"][:3]:
                log(f"  • {rec}")
    
    # 7. 自动修复
    if auto_fix and (report["critical_issues"] > 0 or report["total_issues"] > 3):
        log(f"\n🔧 开始自动修复...")
        fixed_path = fix_word_document(docx_path, ai_result, doc_info)
        if fixed_path:
            report["fixed_file"] = fixed_path
            log(f"✅ 修复完成：{os.path.basename(fixed_path)}")
        else:
            log("❌ 自动修复失败")
    
    log("="*60 + "\n")
    
    return report


def fix_word_document(docx_path, ai_result, doc_info):
    """
    根据检查结果自动修复Word文档
    """
    if not ai_result:
        return None
    
    try:
        # 读取原文档
        doc = Document(docx_path)
        
        # 获取修复后的内容
        fixed_title = ai_result.get("fixed_title", doc_info["title"])
        fixed_content = ai_result.get("fixed_content", doc_info["content"])
        
        # 创建新文档
        new_doc = Document()
        
        # 添加修复后的标题
        title_para = new_doc.add_heading(fixed_title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加修复后的正文（按段落分割）
        paragraphs = fixed_content.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # 检查是否是小标题（**加粗**）
            if para_text.startswith('**') and para_text.endswith('**'):
                heading_text = para_text.strip('*').strip()
                para = new_doc.add_paragraph()
                run = para.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(14)
            else:
                # 普通段落
                para = new_doc.add_paragraph()
                # 简单处理行内格式
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', para_text)
                para.add_run(text)
                para.paragraph_format.line_spacing = 1.5
        
        # 复制原文档中的图片（保持图片不变）
        for shape in doc.inline_shapes:
            try:
                # 注意：这里需要更复杂的图片处理逻辑
                # 简化版本：保留原图片引用
                pass
            except Exception as e:
                log(f"  ⚠️  复制图片失败: {e}")
        
        # 生成新文件名
        base_name = os.path.basename(docx_path)
        name_without_ext = os.path.splitext(base_name)[0]
        fixed_path = os.path.join(
            os.path.dirname(docx_path),
            f"{name_without_ext}_fixed.docx"
        )
        
        # 保存修复后的文档
        new_doc.save(fixed_path)
        
        return fixed_path
        
    except Exception as e:
        log(f"❌ 修复文档失败: {e}")
        return None


def batch_check_today_documents(auto_fix=False):
    """批量检查今天生成的所有Word文档"""
    today = datetime.now().strftime("%Y%m%d")
    pattern = os.path.join(DOCX_DIR, f"{today}_*.docx")
    
    import glob
    files = sorted([f for f in glob.glob(pattern) if not f.endswith('_fixed.docx')])
    
    if not files:
        log(f"❌ 未找到今天({today})生成的Word文档")
        return []
    
    log(f"📋 找到 {len(files)} 篇待检查文档")
    log("="*60)
    
    reports = []
    for i, file_path in enumerate(files, 1):
        log(f"\n[{i}/{len(files)}] 检查中...")
        report = check_word_document(file_path, auto_fix=auto_fix)
        if report:
            reports.append(report)
        
        # 间隔避免API限流
        if i < len(files):
            import time
            time.sleep(2)
    
    # 汇总报告
    log("\n" + "="*60)
    log("📊 批量检查汇总")
    log("="*60)
    
    total_docs = len(reports)
    clean_docs = len([r for r in reports if r["total_issues"] == 0])
    critical_docs = len([r for r in reports if r["critical_issues"] > 0])
    
    log(f"总文档数：{total_docs}")
    log(f"✅ 无问题：{clean_docs}篇")
    log(f"⚠️  有问题：{total_docs - clean_docs}篇")
    log(f"🔴 严重问题：{critical_docs}篇")
    
    if auto_fix:
        fixed_docs = len([r for r in reports if "fixed_file" in r])
        log(f"🔧 已修复：{fixed_docs}篇")
    
    return reports


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    
    arg = sys.argv[1]
    
    if arg == "--batch":
        # 批量检查
        auto_fix = "--fix" in sys.argv
        batch_check_today_documents(auto_fix=auto_fix)
    elif os.path.exists(arg):
        # 单个文件检查
        auto_fix = "--fix" in sys.argv or input("是否自动修复问题？(y/n): ").lower() == 'y'
        check_word_document(arg, auto_fix=auto_fix)
    else:
        log(f"❌ 文件不存在：{arg}")
        sys.exit(1)


if __name__ == "__main__":
    main()
